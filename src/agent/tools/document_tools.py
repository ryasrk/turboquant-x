"""Document generation tools: Word (.docx), PDF, and CSV file creation."""
from __future__ import annotations

import asyncio
import csv
import io
import json
import os
import uuid
from pathlib import Path
from typing import Any

from src.agent.base import Tool

# Generated documents land here; the download endpoint serves from this dir.
_OUTPUT_DIR = Path(os.environ.get("TURBOQUANT_DOC_OUTPUT", "data/generated_docs"))


def _ensure_output_dir() -> Path:
    """Create and return the output directory."""
    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return _OUTPUT_DIR


def _download_link(filename: str) -> str:
    """Return a markdown-formatted download hyperlink.

    Uses a relative URL so it works behind any reverse proxy — the
    browser resolves it against the current page origin.
    """
    return f"[Download {filename}](/v1/documents/download/{filename})"


# ---------------------------------------------------------------------------
# GenerateWordTool
# ---------------------------------------------------------------------------


class GenerateWordTool(Tool):
    """Generate a Microsoft Word (.docx) document and return a download link."""

    @property
    def name(self) -> str:
        return "generate_word"

    @property
    def description(self) -> str:
        return (
            "Create a Word (.docx) document from provided content. "
            "Supports a title, body text (paragraphs separated by newlines), "
            "and optional headers. Returns a download URL for the file."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Document title (appears as heading).",
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Document body. Use '\\n' for paragraph breaks. "
                        "Lines starting with '## ' become sub-headings."
                    ),
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename without extension (default: auto-generated).",
                },
            },
            "required": ["title", "content"],
        }

    async def execute(self, **kwargs: Any) -> str:  # noqa: C901
        title: str = kwargs["title"]
        content: str = kwargs["content"]
        filename: str = kwargs.get("filename") or f"doc_{uuid.uuid4().hex[:8]}"

        # Sanitise filename
        safe_name = "".join(c for c in filename if c.isalnum() or c in ("_", "-"))[:120]
        if not safe_name:
            safe_name = f"doc_{uuid.uuid4().hex[:8]}"

        def _build() -> str:
            from docx import Document  # python-docx
            from docx.shared import Pt

            doc = Document()

            # Title
            doc.add_heading(title, level=0)

            # Body — split into paragraphs; lines starting with "## " become h2
            for para_text in content.split("\n"):
                stripped = para_text.strip()
                if not stripped:
                    doc.add_paragraph("")
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                else:
                    p = doc.add_paragraph(stripped)
                    for run in p.runs:
                        run.font.size = Pt(11)

            out_dir = _ensure_output_dir()
            out_path = out_dir / f"{safe_name}.docx"
            doc.save(str(out_path))
            return str(out_path)

        try:
            path = await asyncio.to_thread(_build)
            fname = os.path.basename(path)
            return (
                f"Word document created: {path}\n"
                f"{_download_link(fname)}"
            )
        except ImportError:
            return "Error: python-docx is not installed. Run: pip install python-docx"
        except Exception as exc:
            return f"Error generating Word document: {exc}"


# ---------------------------------------------------------------------------
# GeneratePdfTool
# ---------------------------------------------------------------------------


class GeneratePdfTool(Tool):
    """Generate a PDF document and return a download link."""

    @property
    def name(self) -> str:
        return "generate_pdf"

    @property
    def description(self) -> str:
        return (
            "Create a PDF document from provided content. "
            "Supports a title, body text (paragraphs separated by newlines), "
            "and optional headers. Returns a download URL for the file."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Document title (appears as heading).",
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Document body. Use '\\n' for paragraph breaks. "
                        "Lines starting with '## ' become sub-headings."
                    ),
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename without extension (default: auto-generated).",
                },
            },
            "required": ["title", "content"],
        }

    async def execute(self, **kwargs: Any) -> str:
        title: str = kwargs["title"]
        content: str = kwargs["content"]
        filename: str = kwargs.get("filename") or f"doc_{uuid.uuid4().hex[:8]}"

        safe_name = "".join(c for c in filename if c.isalnum() or c in ("_", "-"))[:120]
        if not safe_name:
            safe_name = f"doc_{uuid.uuid4().hex[:8]}"

        def _build() -> str:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Title
            pdf.set_font("Helvetica", "B", 20)
            pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(6)

            for para_text in content.split("\n"):
                stripped = para_text.strip()
                if not stripped:
                    pdf.ln(4)
                elif stripped.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 8, stripped[3:], new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)
                elif stripped.startswith("### "):
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.cell(0, 7, stripped[4:], new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)
                else:
                    pdf.set_font("Helvetica", "", 11)
                    pdf.multi_cell(0, 6, stripped)
                    pdf.ln(1)

            out_dir = _ensure_output_dir()
            out_path = out_dir / f"{safe_name}.pdf"
            pdf.output(str(out_path))
            return str(out_path)

        try:
            path = await asyncio.to_thread(_build)
            fname = os.path.basename(path)
            return (
                f"PDF document created: {path}\n"
                f"{_download_link(fname)}"
            )
        except ImportError:
            return "Error: fpdf2 is not installed. Run: pip install fpdf2"
        except Exception as exc:
            return f"Error generating PDF document: {exc}"


# ---------------------------------------------------------------------------
# GenerateCsvTool
# ---------------------------------------------------------------------------


class GenerateCsvTool(Tool):
    """Generate a CSV file and return a download link."""

    @property
    def name(self) -> str:
        return "generate_csv"

    @property
    def description(self) -> str:
        return (
            "Create a CSV file from provided data. "
            "Accepts headers and rows as JSON arrays, or raw CSV text. "
            "Returns a download URL for the file."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "headers": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Column headers (e.g. [\"Name\", \"Age\", \"City\"]).",
                },
                "rows": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "description": 'Rows of data (e.g. [["Alice","30","NYC"],["Bob","25","LA"]]).',
                },
                "raw_csv": {
                    "type": "string",
                    "description": (
                        "Alternative: raw CSV text (including header line). "
                        "Used if headers/rows are not provided."
                    ),
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename without extension (default: auto-generated).",
                },
            },
            "required": [],
        }

    async def execute(self, **kwargs: Any) -> str:
        headers: list[str] | None = kwargs.get("headers")
        rows: list[list[str]] | None = kwargs.get("rows")
        raw_csv: str | None = kwargs.get("raw_csv")
        filename: str = kwargs.get("filename") or f"data_{uuid.uuid4().hex[:8]}"

        safe_name = "".join(c for c in filename if c.isalnum() or c in ("_", "-"))[:120]
        if not safe_name:
            safe_name = f"data_{uuid.uuid4().hex[:8]}"

        if not headers and not rows and not raw_csv:
            return "Error: Provide either (headers + rows) or raw_csv content."

        def _build() -> str:
            out_dir = _ensure_output_dir()
            out_path = out_dir / f"{safe_name}.csv"

            if raw_csv:
                out_path.write_text(raw_csv, encoding="utf-8")
            else:
                with open(out_path, "w", newline="", encoding="utf-8") as fh:
                    writer = csv.writer(fh)
                    if headers:
                        writer.writerow(headers)
                    if rows:
                        writer.writerows(rows)

            return str(out_path)

        try:
            path = await asyncio.to_thread(_build)
            fname = os.path.basename(path)
            return (
                f"CSV file created: {path}\n"
                f"{_download_link(fname)}"
            )
        except Exception as exc:
            return f"Error generating CSV file: {exc}"
